from pyresparser import ResumeParser
import numpy as np
import time
import json
from pdfTextExtract import extract_text_from_pdf
from docx import Document
from find_job_titles import FinderAcora
skillsData = np.load("finalDataTags.npy")


import docx

'''
Method to check if education heading exists in a string.
Add more values to list to increase search criteria
'''
def checkEducation(txt):
    edu = ["Education", "Degree"]
    for e in edu:
        if e.lower() in txt.lower():
            return True
    return False


'''
Method to check if Experience heading exists in a string.
Add more values to list to increase search criteria
'''
def checkExperience(txt):
    exp = ["Experience", "Work Experience"]
    for e in exp:
        if e.lower() in txt.lower():
            return True
    return False


# Get text from a docx file
def getTextFromDocx(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)
    
'''
Finalize meta results. Works on resume and job description.
Check headings in resume and jobs
Find job title in job text. Match job title with resume
Check length of resume
'''
def finalizeMetaResults(ResumePath, JobPath, ResumeData, JobData):
    meta = {}
    text = getTextFromDocx(ResumePath)
    meta["length"] = len(text)
    textJob = getTextFromDocx(JobPath)
    finder = FinderAcora()
    jobTitle = finder.findall(textJob)
    if len(jobTitle) != 0:
        jobTitle = jobTitle[0].match
        meta["jobTitle"] = jobTitle
        meta["jobTitleResume"] = "notMatch"
        if jobTitle in text:
            meta["jobTitleResume"] = "match"
    
    
    meta["educationHeading"] = True if checkEducation(text) else False    
    meta["workExperienceHeading"] = True if checkExperience(text) else False
    meta["educationRequirements"] = True if checkEducation(textJob) else False
    ResumeData["skillsData"] = countWords(ResumeData["skillsData"], text.lower(), textJob.lower())
    
    return meta

'''
Count number of occurence of Skills in resume and job
'''
def countWords(skills, ResumeText, JobText):
    meta = {}
    for skill, val in skills.items():
        skillTest = skill.lower()
        meta[skill] = {
            "ResumeCount" : ResumeText.count(skillTest),
            "JobCount" : JobText.count(skillTest),
            "wanted" : val["wanted"],
            "exists" : val["exists"]
        }
    return meta

'''
Cross validation for skills to check if skill name is valid
'''
def parseSkills(skills):
    skillsFinal = []
    for skill in skills:
        lowerSkill = skill.lower()
        if lowerSkill in skillsData:
            skillsFinal.append(skill)
    return skillsFinal



'''
    Method to calculate ATS attributes score.

'''
def calcAtsScore(resumeData, fileName, skills):
    points = {}
    points["skills"] = 1 if skills["not-exists"] == 0 else 0
    print(resumeData["experience"])
    
    if resumeData["experience"] != None:
        points["experience"] = 1 if len(resumeData["experience"]) > 0 else 0
    else:
        points["experience"] = 0
    points["email"] = 1 if resumeData["email"] is not None else 0
        
    points["mobile_number"] = 1 if resumeData["mobile_number"] is not None else 0
         
    points["name"] = 1 if resumeData["name"] is not None else 0
          
    points["jobTitleMatch"] = 1 if resumeData["meta"]["jobTitleResume"] == "exists" else 0
        
    if resumeData["meta"]["educationRequirements"]:
        if resumeData["meta"]["educationHeading"]:
            points["educationReq"] = 1
    else:
        points["educationReq"] = 1
    points["educationHeading"] = 1 if resumeData["meta"]["educationHeading"] else 0
    points["workHeading"] = 1 if resumeData["meta"]["workExperienceHeading"] else 0
    if fileName != None:
        if fileName.split(".")[-1] in ["pdf", "docx"]:
            points["filePoint"] = 1
            points["fileNamePoint"] = 1
            for i in ["-", "@", "!", "$", "^", "&", "*"]:
                if i in fileName:
                    points["fileNamePoint"] = 0
            if points["fileNamePoint"] == 1:
                points["readAble"] = 1

    totalScore = 0
    totalItems = len(points.keys())
    for val in points.values():
        totalScore += val
    p = {
        "exists": totalScore,
        "not-exists": totalItems - totalScore,
        "total": totalItems
    }
    return p

'''
    Method to count Recruiter score.
'''

def checkRecruiterScore(resumeData):
    points = {}
    points["wordCount"] = 1 if resumeData["meta"]["length"] < 1000 else 0
    points["measureableResults"] = 1
    points["avoidWords"] = 0
    points["jobLevel"] = 1
    i, total = 0, 0
    for val in points.values():
        total += val
    p = {
        "exists": total,
        "not-exists": 4 - total,
        "total": 4
    }
    return p


'''
    Check skills score
'''
def checkSkillsScore(resumeSkills, JobSkills):
    resumeSkills = [skill.lower() for skill in resumeSkills]
    JobSkills = [skill.lower() for skill in JobSkills]
    eScore = 0
    for jSkill in JobSkills:
        if jSkill in resumeSkills:
            eScore += 1
            
    return {"exists": eScore, "not-exists" : len(JobSkills) - eScore, "total" : len(JobSkills)}

'''
    Method wrapper for all points count for
    skills
    ats
    and recruiter findings

'''

def calculatePoints(resumeData, jobData, fileName):
    skills = checkSkillsScore(resumeData["skills"], jobData["skills"])
    data = {
        "skills": skills,
        "ats": calcAtsScore(resumeData, fileName, skills),
        "rfindings": checkRecruiterScore(resumeData)
    }
    return data


'''

    Calculate Final match score
    Criteria:
        Skills: 50%
        ATS: 25%
        RecruiterFindings: 25%
'''
def calculateMatch(score):
    skillsMatchScore = (score["skills"]["exists"] / score["skills"]["total"]) * 0.5
    atsMatchScore = (score["ats"]["exists"] / score["ats"]["total"]) * 0.25
    rfindingsScore = (score["rfindings"]["exists"] / score["rfindings"]["total"]) * 0.25
    totalScore = skillsMatchScore + atsMatchScore + rfindingsScore
    return round(totalScore * 100)


'''
    Format data in special format to send back to client

'''
def finalArrangeData(resumeData, fileName):
    data = {}
    data["totalScore"] = resumeData["matchScore"]
    data["scores"] = resumeData["scores"]
    data["ats"] = {
        "name" : resumeData["name"],
        "email" : resumeData["email"],
        "mobileNumber" : resumeData["mobile_number"],
        "resumeSkillsMissing" : resumeData["scores"]["skills"]["not-exists"],
        "jobTitleMatch" : resumeData["meta"]["jobTitleResume"],
        "jobTitle" : resumeData["meta"]["jobTitle"],
        "educationMatch" : {"required": resumeData["meta"]["educationRequirements"], "match": False},
        "headings" : {"educationHeading" : resumeData["meta"]["educationHeading"], "workExperienceHeading" : resumeData["meta"]["workExperienceHeading"]},
        "dateFormatting" : True
    }
    if fileName != None:
        if fileName.split(".")[-1] in ["pdf", "docx"]:
            noSpecialCharName = True
            for i in ["-", "@", "!", "$", "^", "&", "*"]:
                if i in fileName:
                    noSpecialCharName = False
            data["ats"]["fileFormat"] = {
                "format" : fileName.split(".")[-1],
                "noSpecialCharName" : noSpecialCharName,
                "nameReadable" : True
            }
    data["recruiterFindings"] = {
        "length" : {"current" : resumeData["meta"]["length"], "allowed" : 1000},
        "measureableResults" : None,
        "wordsToAvoid" : None,
        "jobLevelMatch" : True
    }
    data["skillsData"] = resumeData["skillsData"]
    data["meta"] = {
        "noOfPages" : resumeData["no_of_pages"]
    }
    return data



'''
Finalize skills for display. 
Checks if skill exists in resume or job

'''
def finalizeSkillsDisplay(resumeSkills, JobSkills):
    fSkills = {}
    resumeSkills = [skill.lower() for skill in resumeSkills]

    JobSkills = [skill.lower() for skill in JobSkills]
    eSkills = 0
    for skill in resumeSkills:
        if skill not in JobSkills:
            fSkills[skill] = {
                "wanted": False,
                "exists": True
            }
        else:
            eSkills += 1
            fSkills[skill] = {
                "wanted": True,
                "exists": True
            }
    for skill in JobSkills:
        if skill not in resumeSkills:
            fSkills[skill] = {
                "wanted": True,
                "exists": False
            }
        else:
            fSkills[skill] = {
                "wanted": True,
                "exists": False
            }
    return fSkills


# Wrapper funcation for get raw results for all data
def getResults(resumePath, jobPath):
    old = time.time()
    resumeData = ResumeParser(resumePath).get_extracted_data()
    resumeData["skills"] = parseSkills(resumeData["skills"])
    #print("Resume Skills:", resumeData["skills"])

    jobData = ResumeParser(jobPath).get_extracted_data()
    jobData["skills"] = parseSkills(jobData["skills"])
    #print("Job Skills:", jobData["skills"])

    skillsRes = finalizeSkillsDisplay(resumeData["skills"], jobData["skills"])
    resumeData["skillsData"] = skillsRes

    miscResults = finalizeMetaResults(resumePath, jobPath, resumeData, jobData)

    resumeData["meta"] = miscResults

    return resumeData, jobData


# main funcation.
'''
Send resume path, job description path and fileType

'''
def parseAndMatchResume(resumePath, jobPath, fileName):
    rawResultsResume, rawResultsJob = getResults(resumePath, jobPath)
    rawResultsResume["scores"] = calculatePoints(rawResultsResume, rawResultsJob, fileName)
    rawResultsResume["matchScore"] = calculateMatch(rawResultsResume["scores"])
    finalResults = finalArrangeData(rawResultsResume, fileName)
    return finalResults
    

# Validate data
def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
        )


def CleanString(text):
    cleaned_string = ''.join(c for c in text if valid_xml_char_ordinal(c))
    return cleaned_string

# Creates a docx file

def buildDocxFile(data, filePath):
    data = CleanString(data)
    document = Document()
    p = document.add_paragraph(data)
    document.save(filePath)


# Call function to process text form route.
# To be called from route
def processText(resumeText, jobText, ftype):
    resumePath = "resume.docx"
    jobPath = "job.docx"
    buildDocxFile(resumeText, resumePath)
    buildDocxFile(jobText, jobPath)
    fileType = None if ftype == "text" else "file." + ftype
    res = parseAndMatchResume(resumePath, jobPath, fileType)
    print(json.dumps(res, indent=2))


# Get text from file
# Allowed : PDF of DOCX
def convertFileToText(fileName):
    ext = fileName.split(".")[-1]
    if ext == "pdf":
        text = ""
        for page in extract_text_from_pdf(ResumePath):
            text += ' ' + page
        return text
    elif ext == "docx":
        text = getTextFromDocx(fileName)
        return text
    else:
        None




#text = ""
#for page in extract_text_from_pdf("C:/Users/Faizan/Downloads/Microsoft.SkypeApp_kzf8qxf38zg5c!App/All/Profile (1).pdf"):
#    text += ' ' + page
#cleaned_string = ''.join(c for c in text if valid_xml_char_ordinal(c))
#buildDocxFile(cleaned_string, "resume.docx")
#data = parseAndMatchResume("resume.docx","demo.docx", "profile.pdf")
#print(data)